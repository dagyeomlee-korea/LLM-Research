"""Qwen3.5 10개 모델 계보 분석 팀 공유 DOCX 생성기."""

from __future__ import annotations

import argparse
import math
import zipfile
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


NAVY = "17365D"
BLUE = "DCE6F1"
LIGHT_BLUE = "EEF4FB"
GREEN = "E2F0D9"
YELLOW = "FFF2CC"
ORANGE = "FCE4D6"
GRAY = "E7E6E6"
LIGHT_GRAY = "F5F5F5"
RED = "F4CCCC"
WHITE = "FFFFFF"


MODEL_ROWS = [
    (
        "1",
        "Qwen/Qwen3.5-0.8B-Base",
        "0.8B 공식 Base",
        "공식 기준점",
        "Instruct와 core 전반에서 차이",
        "기준점",
    ),
    (
        "2",
        "Qwen/Qwen3.5-0.8B",
        "0.8B 공식 Instruct",
        "공식 기준점",
        "0.8B 파생 후보의 최근접 anchor",
        "기준점",
    ),
    (
        "3",
        "unsloth/Qwen3.5-0.8B",
        "0.8B Instruct",
        "Exact mirror",
        "SHA-256 전체 weight 파일 동일",
        "확정",
    ),
    (
        "4",
        "CloudGoat/Qwen3.5-0.8B-JP-Tuned-v1.0",
        "0.8B Instruct",
        "구조 보존형 광범위 core 조정 후보",
        "최근접, 96개 core matrix 변경, CKA 0.991 이상",
        "강한 후보",
    ),
    (
        "5",
        "huihui-ai/Huihui-Qwen3.5-0.8B-abliterated",
        "0.8B Instruct",
        "표적 rank-1 projection 변환 후보",
        "O/down 집중, top-1 energy 99.71% 이상",
        "강한 후보",
    ),
    (
        "6",
        "Qwen/Qwen3.5-2B-Base",
        "2B 공식 Base",
        "공식 기준점",
        "Instruct와 core 전반에서 차이",
        "기준점",
    ),
    (
        "7",
        "Qwen/Qwen3.5-2B",
        "2B 공식 Instruct",
        "공식 기준점",
        "2B 파생 후보의 최근접 anchor",
        "기준점",
    ),
    (
        "8",
        "unsloth/Qwen3.5-2B",
        "2B Instruct",
        "Exact mirror",
        "SHA-256 전체 weight 파일 동일",
        "확정",
    ),
    (
        "9",
        "hamishivi/Qwen3.5-2B",
        "2B Instruct",
        "Exact mirror",
        "SHA-256 전체 weight 파일 동일",
        "확정",
    ),
    (
        "10",
        "huihui-ai/Huihui-Qwen3.5-2B-abliterated",
        "2B Instruct",
        "표적 rank-1 projection 변환 후보",
        "O/down 집중, top-1 energy 99.46% 이상",
        "강한 후보",
    ),
]


def parse_args() -> argparse.Namespace:
    """
    목적: DOCX 생성에 필요한 출력 경로와 한글 글꼴 경로를 수집한다.
    입력: 명령행의 --output-dir, --font, --bold-font 옵션.
    처리: argparse로 필수·선택 인자를 검증하고 Path 변환 전 원문을 보존한다.
    반환/부작용: argparse.Namespace를 반환하며 잘못된 인자이면 사용법을 출력한다.
    """

    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--font", type=Path, required=True)
    parser.add_argument("--bold-font", type=Path)
    return parser.parse_args()


def load_font(path: Path, size: int) -> ImageFont.FreeTypeFont:
    """
    목적: PNG 도식에서 한글을 정상 렌더링할 TrueType 글꼴을 연다.
    입력: 글꼴 파일 경로와 픽셀 단위 크기.
    처리: 파일 존재 여부를 확인하고 Pillow FreeTypeFont로 로드한다.
    반환/부작용: 로드된 글꼴을 반환하며 파일이 없으면 예외를 발생시킨다.
    """

    if not path.is_file():
        raise FileNotFoundError(f"글꼴 파일을 찾을 수 없습니다: {path}")
    return ImageFont.truetype(str(path), size=size)


def wrap_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    font: ImageFont.FreeTypeFont,
    max_width: int,
) -> list[str]:
    """
    목적: 도식 상자 안에 한글·영문 문장을 폭 기준으로 줄바꿈한다.
    입력: Pillow Draw 객체, 원문, 글꼴, 허용 픽셀 폭.
    처리: 공백 단위 토큰을 누적하고 폭을 넘으면 새 줄을 시작한다.
    반환/부작용: 줄 문자열 목록을 반환하며 이미지에는 직접 그리지 않는다.
    """

    lines: list[str] = []
    for source_line in text.splitlines() or [""]:
        words = source_line.split(" ")
        current = ""
        for word in words:
            candidate = word if not current else f"{current} {word}"
            width = draw.textbbox((0, 0), candidate, font=font)[2]
            if current and width > max_width:
                lines.append(current)
                current = word
            else:
                current = candidate
        lines.append(current)
    return lines


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: str,
    spacing: int = 8,
) -> None:
    """
    목적: 지정한 사각형 중앙에 자동 줄바꿈된 텍스트를 배치한다.
    입력: Draw 객체, 좌상·우하 좌표, 문구, 글꼴, 색상, 행간.
    처리: 폭에 맞춰 문장을 나누고 전체 높이를 계산해 수직·수평 중앙 정렬한다.
    반환/부작용: 반환값 없이 대상 이미지에 텍스트를 그린다.
    """

    left, top, right, bottom = box
    lines = wrap_text(draw, text, font, max_width=(right - left) - 36)
    metrics = [draw.textbbox((0, 0), line, font=font) for line in lines]
    heights = [metric[3] - metric[1] for metric in metrics]
    total_height = sum(heights) + spacing * max(0, len(lines) - 1)
    y = top + ((bottom - top) - total_height) / 2
    for line, metric, line_height in zip(lines, metrics, heights):
        line_width = metric[2] - metric[0]
        x = left + ((right - left) - line_width) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height + spacing


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: str,
    outline: str,
    text_fill: str = "#1F1F1F",
    width: int = 4,
) -> None:
    """
    목적: 계보·파이프라인 노드를 둥근 사각형으로 표현한다.
    입력: Draw 객체, 좌표, 문구, 글꼴, 배경·테두리·글자색과 선 두께.
    처리: rounded_rectangle을 그린 뒤 중앙 정렬 텍스트를 삽입한다.
    반환/부작용: 반환값 없이 대상 이미지 픽셀을 변경한다.
    """

    draw.rounded_rectangle(box, radius=22, fill=fill, outline=outline, width=width)
    draw_centered_text(draw, box, text, font, text_fill)


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: str,
    width: int = 5,
    dashed: bool = False,
    arrow: bool = True,
) -> None:
    """
    목적: 노드 사이의 확정 관계·후보 방향을 선과 화살표로 구분한다.
    입력: Draw 객체, 시작·끝 좌표, 색, 두께, 점선 여부, 화살촉 여부.
    처리: 점선이면 구간별로 선을 그리고 요청 시 끝점에 삼각 화살촉을 추가한다.
    반환/부작용: 반환값 없이 대상 이미지에 관계선을 그린다.
    """

    sx, sy = start
    ex, ey = end
    if dashed:
        dx, dy = ex - sx, ey - sy
        distance = max(1.0, math.hypot(dx, dy))
        ux, uy = dx / distance, dy / distance
        cursor = 0.0
        while cursor < distance - 18:
            seg_end = min(cursor + 18, distance - 18)
            draw.line(
                (
                    sx + ux * cursor,
                    sy + uy * cursor,
                    sx + ux * seg_end,
                    sy + uy * seg_end,
                ),
                fill=color,
                width=width,
            )
            cursor += 30
    else:
        draw.line((sx, sy, ex, ey), fill=color, width=width)

    if arrow:
        angle = math.atan2(ey - sy, ex - sx)
        length = 18
        spread = math.pi / 6
        points = [
            (ex, ey),
            (
                ex - length * math.cos(angle - spread),
                ey - length * math.sin(angle - spread),
            ),
            (
                ex - length * math.cos(angle + spread),
                ey - length * math.sin(angle + spread),
            ),
        ]
        draw.polygon(points, fill=color)


def create_pipeline_image(output_path: Path, font_path: Path, bold_path: Path) -> None:
    """
    목적: 10개 weight에서 MotherTree 후보도까지 이어지는 분석 파이프라인을 PNG로 만든다.
    입력: 출력 PNG 경로와 일반·굵은 한글 글꼴 경로.
    처리: 단계 노드, 화살표, all/core 분기와 지표 역할을 고해상도 캔버스에 그린다.
    반환/부작용: 반환값 없이 output_path에 PNG 파일을 생성한다.
    """

    output_path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1900, 1040), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(bold_path, 50)
    node_font = load_font(bold_path, 30)
    body_font = load_font(font_path, 25)
    small_font = load_font(font_path, 22)

    draw.text((70, 45), "Qwen3.5 가중치 계보 분석 파이프라인", font=title_font, fill="#17365D")
    nodes = [
        ((70, 170, 350, 320), "1. Weight 수집\n10개 / 29.35GB"),
        ((440, 170, 720, 320), "2. SHA-256\n완전 동일 그룹"),
        ((810, 170, 1090, 320), "3. Family 분리\n0.8B / 2B"),
        ((1180, 170, 1460, 320), "4. Tensor 정렬\n이름·shape·dtype"),
        ((1550, 170, 1830, 320), "5. 거리·분포\nL2 / cosine / kurtosis"),
    ]
    for box, label in nodes:
        draw_box(draw, box, label, node_font, "#EEF4FB", "#2457A6")
    for idx in range(len(nodes) - 1):
        draw_arrow(
            draw,
            (nodes[idx][0][2] + 12, 245),
            (nodes[idx + 1][0][0] - 12, 245),
            "#2457A6",
        )

    lower_nodes = [
        ((250, 520, 650, 700), "6A. 광범위 변경\nCKA로 geometry 보존 확인"),
        ((750, 520, 1150, 700), "6B. 국소 변경\nDelta SVD로 rank 확인"),
        ((1250, 520, 1650, 700), "7. MotherTree\n거리=edge 비용\nCKA/SVD=edge 유형"),
    ]
    draw_box(draw, lower_nodes[0][0], lower_nodes[0][1], node_font, "#FFF2CC", "#A56B00")
    draw_box(draw, lower_nodes[1][0], lower_nodes[1][1], node_font, "#FCE4D6", "#C65911")
    draw_box(draw, lower_nodes[2][0], lower_nodes[2][1], node_font, "#E2F0D9", "#1F7A3D")
    draw_arrow(draw, (1690, 340), (490, 500), "#A56B00", dashed=True)
    draw_arrow(draw, (1690, 340), (950, 500), "#C65911", dashed=True)
    # 6A 연결선은 6B 상자를 통과하지 않도록 위쪽으로 우회한다.
    draw_arrow(draw, (620, 520), (1180, 420), "#1F7A3D", arrow=False)
    draw_arrow(draw, (1180, 420), (1230, 575), "#1F7A3D")
    draw_arrow(draw, (1150, 650), (1230, 650), "#1F7A3D")

    draw.rounded_rectangle((120, 820, 1780, 965), radius=18, fill="#F5F5F5", outline="#A6A6A6", width=3)
    draw.text((160, 845), "핵심 원칙", font=node_font, fill="#17365D")
    explanation = (
        "SHA는 '완전히 같은가', weight distance는 '얼마나 움직였나', kurtosis는 '분포 모양이 달라졌나', "
        "CKA는 '관계 구조가 유지됐나', SVD는 '변화가 몇 개 방향인가'에 답한다."
    )
    lines = wrap_text(draw, explanation, body_font, 1420)
    for index, line in enumerate(lines):
        draw.text((390, 848 + index * 38), line, font=body_font, fill="#333333")
    draw.text((1510, 985), "실측 결과 기반", font=small_font, fill="#666666")
    image.save(output_path, format="PNG", dpi=(180, 180))


def create_mothertree_image(output_path: Path, font_path: Path, bold_path: Path) -> None:
    """
    목적: 10개 모델을 모두 포함하고 증거 수준을 구분한 MotherTree 후보도를 만든다.
    입력: 출력 PNG 경로와 일반·굵은 한글 글꼴 경로.
    처리: 0.8B·2B 패널에 공식 anchor, exact mirror, 파생 후보와 관계선을 그린다.
    반환/부작용: 반환값 없이 output_path에 고해상도 PNG를 생성한다.
    """

    output_path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (2000, 1360), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(bold_path, 52)
    panel_font = load_font(bold_path, 36)
    node_font = load_font(bold_path, 27)
    label_font = load_font(font_path, 23)
    legend_font = load_font(font_path, 24)

    draw.text((70, 35), "Qwen3.5 10개 모델 MotherTree 후보도", font=title_font, fill="#17365D")
    panels = [(50, 125, 1950, 650), (50, 690, 1950, 1215)]
    for panel in panels:
        draw.rounded_rectangle(panel, radius=26, fill="#FAFAFA", outline="#BFBFBF", width=3)
    draw.text((90, 150), "0.8B family — 배포본 5개 / 고유 weight 4개", font=panel_font, fill="#17365D")
    draw.text((90, 715), "2B family — 배포본 5개 / 고유 weight 3개", font=panel_font, fill="#17365D")

    # 0.8B family
    b08 = (90, 295, 400, 425)
    i08 = (550, 295, 900, 425)
    u08 = (1050, 210, 1390, 340)
    c08 = (1050, 375, 1450, 520)
    h08 = (1560, 375, 1910, 520)
    draw_box(draw, b08, "Qwen 0.8B Base\n공식 anchor", node_font, "#DCE6F1", "#2457A6")
    draw_box(draw, i08, "Qwen 0.8B Instruct\n공식 anchor", node_font, "#DCE6F1", "#2457A6")
    draw_box(draw, u08, "unsloth 0.8B\nExact mirror", node_font, "#E2F0D9", "#1F7A3D")
    draw_box(draw, c08, "CloudGoat 0.8B\n광범위 core 조정 후보", node_font, "#FFF2CC", "#A56B00")
    draw_box(draw, h08, "Huihui 0.8B\nO/down rank-1 후보", node_font, "#FCE4D6", "#C65911")
    draw_arrow(draw, (420, 360), (530, 360), "#7F7F7F", dashed=True)
    draw.text((414, 325), "참고 방향", font=label_font, fill="#666666")
    draw_arrow(draw, (920, 335), (1030, 285), "#1F7A3D", arrow=False, width=8)
    draw.text((915, 245), "SHA 동일", font=label_font, fill="#1F7A3D")
    draw_arrow(draw, (920, 385), (1030, 440), "#A56B00", dashed=True)
    draw.text((930, 420), "최근접 + CKA", font=label_font, fill="#A56B00")
    # Huihui 연결선은 CloudGoat 상자를 가로지르지 않도록 상단으로 우회한다.
    draw_arrow(draw, (920, 400), (1450, 345), "#C65911", dashed=True, arrow=False)
    draw_arrow(draw, (1450, 345), (1540, 430), "#C65911", dashed=True)
    draw.text((1240, 310), "최근접 + SVD", font=label_font, fill="#C65911")

    # 2B family
    b2 = (90, 860, 400, 990)
    i2 = (550, 860, 900, 990)
    u2 = (1050, 765, 1370, 885)
    m2 = (1050, 975, 1370, 1095)
    h2 = (1540, 860, 1910, 1005)
    draw_box(draw, b2, "Qwen 2B Base\n공식 anchor", node_font, "#DCE6F1", "#2457A6")
    draw_box(draw, i2, "Qwen 2B Instruct\n공식 anchor", node_font, "#DCE6F1", "#2457A6")
    draw_box(draw, u2, "unsloth 2B\nExact mirror", node_font, "#E2F0D9", "#1F7A3D")
    draw_box(draw, m2, "hamishivi 2B\nExact mirror", node_font, "#E2F0D9", "#1F7A3D")
    draw_box(draw, h2, "Huihui 2B\nO/down rank-1 후보", node_font, "#FCE4D6", "#C65911")
    draw_arrow(draw, (420, 925), (530, 925), "#7F7F7F", dashed=True)
    draw.text((414, 890), "참고 방향", font=label_font, fill="#666666")
    draw_arrow(draw, (920, 890), (1030, 830), "#1F7A3D", arrow=False, width=8)
    draw_arrow(draw, (920, 960), (1030, 1035), "#1F7A3D", arrow=False, width=8)
    draw.text((925, 800), "SHA 동일", font=label_font, fill="#1F7A3D")
    draw.text((925, 1055), "SHA 동일", font=label_font, fill="#1F7A3D")
    draw_arrow(draw, (920, 930), (1520, 930), "#C65911", dashed=True)
    draw.text((1230, 890), "최근접 + SVD", font=label_font, fill="#C65911")

    draw.text((90, 1260), "범례:", font=panel_font, fill="#17365D")
    draw.line((260, 1282, 390, 1282), fill="#1F7A3D", width=9)
    draw.text((410, 1260), "SHA-256 동일(확정)", font=legend_font, fill="#333333")
    draw_arrow(draw, (740, 1282), (870, 1282), "#C65911", dashed=True)
    draw.text((890, 1260), "가중치 기반 파생 후보", font=legend_font, fill="#333333")
    draw_arrow(draw, (1300, 1282), (1430, 1282), "#7F7F7F", dashed=True)
    draw.text((1450, 1260), "시간 방향 미확정", font=legend_font, fill="#333333")
    image.save(output_path, format="PNG", dpi=(180, 180))


def set_cell_shading(cell, fill: str) -> None:
    """
    목적: Word 표 셀에 증거 수준과 헤더를 구분하는 배경색을 적용한다.
    입력: python-docx 셀 객체와 6자리 RGB 색상 문자열.
    처리: 셀 속성 XML에 w:shd 요소를 추가하거나 기존 값을 갱신한다.
    반환/부작용: 반환값 없이 셀 OpenXML을 변경한다.
    """

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 90, bottom: int = 90, end: int = 90) -> None:
    """
    목적: 표 셀 내부 여백을 통일해 긴 한글 표의 가독성을 높인다.
    입력: 셀 객체와 twip 단위 상·좌·하·우 여백.
    처리: tcMar XML 요소와 네 방향 값을 생성·갱신한다.
    반환/부작용: 반환값 없이 셀 OpenXML 속성을 변경한다.
    """

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    """
    목적: 여러 페이지로 넘어가는 표에서 첫 행을 반복 헤더로 지정한다.
    입력: 헤더로 사용할 python-docx Row 객체.
    처리: 행 속성에 tblHeader OpenXML 요소를 추가한다.
    반환/부작용: 반환값 없이 행 속성을 변경한다.
    """

    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_cell_text(cell, text: str, bold: bool = False, size: float = 8.0) -> None:
    """
    목적: 표 셀의 기존 내용을 지우고 통일된 글꼴·크기로 텍스트를 입력한다.
    입력: 셀, 문자열, 굵게 여부, 포인트 크기.
    처리: 첫 문단을 재사용하고 run을 생성해 맑은 고딕 글꼴을 지정한다.
    반환/부작용: 반환값 없이 셀 텍스트와 서식을 변경한다.
    """

    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_table(
    document: Document,
    headers: Sequence[str],
    rows: Iterable[Sequence[str]],
    widths_cm: Sequence[float] | None = None,
    font_size: float = 8.5,
) -> object:
    """
    목적: 헤더 색상과 반복 행이 적용된 표준 보고서 표를 추가한다.
    입력: 문서, 헤더, 행 데이터, 선택적 열 너비, 본문 글꼴 크기.
    처리: Table Grid 표를 만들고 헤더·본문 셀 서식과 증거 수준 색상을 적용한다.
    반환/부작용: 생성한 Table 객체를 반환하고 문서에 표를 삽입한다.
    """

    row_values = list(rows)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = widths_cm is None
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_repeat_cell_text(cell, header, bold=True, size=8.5)
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
        if widths_cm:
            cell.width = Cm(widths_cm[index])
    set_repeat_table_header(table.rows[0])

    for values in row_values:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_repeat_cell_text(cells[index], str(value), size=font_size)
            if widths_cm:
                cells[index].width = Cm(widths_cm[index])
        evidence = str(values[-1]) if values else ""
        if evidence == "확정":
            set_cell_shading(cells[-1], GREEN)
        elif evidence == "강한 후보":
            set_cell_shading(cells[-1], YELLOW)
        elif evidence == "미확정":
            set_cell_shading(cells[-1], RED)
        elif evidence == "기준점":
            set_cell_shading(cells[-1], BLUE)
    document.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def set_document_font(document: Document) -> None:
    """
    목적: DOCX 전체 기본 글꼴과 제목 계층을 한글 보고서용으로 통일한다.
    입력: python-docx Document 객체.
    처리: Normal·Title·Heading 스타일의 글꼴, 색, 크기, 간격을 설정한다.
    반환/부작용: 반환값 없이 문서 스타일 컬렉션을 변경한다.
    """

    normal = document.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    styles = [
        ("Title", 25, NAVY, 0),
        ("Subtitle", 12, "666666", 0),
        ("Heading 1", 17, NAVY, 10),
        ("Heading 2", 13.5, "2457A6", 7),
        ("Heading 3", 11.5, "7F6000", 5),
    ]
    for style_name, size, color, before in styles:
        style = document.styles[style_name]
        style.font.name = "맑은 고딕"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(5)


def set_page_layout(document: Document) -> None:
    """
    목적: 문서를 A4 세로형 팀 보고서 레이아웃으로 설정한다.
    입력: Document 객체.
    처리: 모든 section의 용지, 여백, 머리글·바닥글 거리를 지정한다.
    반환/부작용: 반환값 없이 section 레이아웃을 변경한다.
    """

    for section in document.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.7)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.8)


def add_page_number(paragraph) -> None:
    """
    목적: 바닥글 문단에 Word가 갱신할 수 있는 PAGE 필드를 삽입한다.
    입력: 바닥글의 python-docx Paragraph 객체.
    처리: 필드 시작·명령·종료 OpenXML run을 순서대로 추가한다.
    반환/부작용: 반환값 없이 문단에 동적 페이지 번호 필드를 추가한다.
    """

    run = paragraph.add_run("페이지 ")
    run.font.name = "맑은 고딕"
    run.font.size = Pt(8)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def add_header_footer(document: Document) -> None:
    """
    목적: 보고서 식별용 머리글과 페이지 번호 바닥글을 추가한다.
    입력: Document 객체.
    처리: 첫 section의 header/footer 문단에 문서명과 PAGE 필드를 배치한다.
    반환/부작용: 반환값 없이 머리글·바닥글 내용을 변경한다.
    """

    for section in document.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run("Qwen3.5 10개 모델 가중치 계보 분석")
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(127, 127, 127)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(footer)


def add_bullets(document: Document, items: Iterable[str], level: int = 0) -> None:
    """
    목적: 핵심 결과와 설명 절차를 일관된 글머리표 목록으로 추가한다.
    입력: Document 객체, 문장 iterable, 들여쓰기 수준.
    처리: List Bullet 스타일 문단을 만들고 level에 따라 왼쪽 들여쓰기를 설정한다.
    반환/부작용: 반환값 없이 문서 끝에 목록 문단을 추가한다.
    """

    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
        paragraph.paragraph_format.first_line_indent = Cm(-0.25)
        paragraph.add_run(item)


def add_numbered(document: Document, items: Iterable[str]) -> None:
    """
    목적: 분석 알고리즘 순서를 번호 목록으로 문서에 추가한다.
    입력: Document 객체와 순서가 있는 문장 iterable.
    처리: List Number 스타일 문단을 생성한다.
    반환/부작용: 반환값 없이 문서 끝에 번호 목록을 추가한다.
    """

    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Cm(0.7)
        paragraph.paragraph_format.first_line_indent = Cm(-0.3)
        paragraph.add_run(item)


def add_callout(document: Document, title: str, body: str, fill: str = LIGHT_BLUE) -> None:
    """
    목적: 결론·주의사항을 본문과 구별되는 한 칸 표 형태로 강조한다.
    입력: Document 객체, 강조 제목, 본문, 배경색.
    처리: 1열 1행 표를 만들고 굵은 제목과 일반 본문을 같은 셀에 배치한다.
    반환/부작용: 반환값 없이 문서에 강조 상자를 추가한다.
    """

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    title_run = paragraph.add_run(f"{title}\n")
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(NAVY)
    body_run = paragraph.add_run(body)
    body_run.font.size = Pt(10.5)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_metric_section(
    document: Document,
    heading: str,
    easy: str,
    principle: str,
    role: str,
    caution: str,
) -> None:
    """
    목적: 전문 지표를 초심자에게 동일한 네 단계 형식으로 설명한다.
    입력: 문서, 지표 제목, 쉬운 설명, 원리, 분석 역할, 주의점.
    처리: Heading 2와 2열 정의 표를 추가해 개념별 설명 구조를 통일한다.
    반환/부작용: 반환값 없이 문서에 지표 설명 절을 추가한다.
    """

    document.add_heading(heading, level=2)
    add_table(
        document,
        ["구분", "설명"],
        [
            ("쉬운 설명", easy),
            ("원리", principle),
            ("이번 분석의 역할", role),
            ("주의", caution),
        ],
        widths_cm=[3.2, 13.6],
        font_size=9.3,
    )


def add_title_page(document: Document) -> None:
    """
    목적: 팀 공유 DOCX의 제목·대상·핵심 결론이 보이는 표지를 만든다.
    입력: Document 객체.
    처리: 상단 여백, 제목, 부제, 날짜, 핵심 메시지 강조 상자를 배치한다.
    반환/부작용: 반환값 없이 표지와 페이지 나눔을 문서에 추가한다.
    """

    document.add_paragraph().paragraph_format.space_after = Pt(60)
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Qwen3.5 10개 모델\n가중치 계보 분석")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("초심자용 팀 공유 보고서 · MotherTree · 알고리즘 원리")
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("분석 기준일: 2026-07-13\n대상: Qwen3.5 0.8B 5개 + 2B 5개")
    document.add_paragraph().paragraph_format.space_after = Pt(35)
    add_callout(
        document,
        "결론 한 줄",
        "10개 배포본은 고유 weight 7개로 정리된다. 3개는 공식 Instruct와 byte 단위로 동일하고, "
        "CloudGoat는 구조 보존형 광범위 조정 후보, Huihui 0.8B/2B는 O/down의 표적 rank-1 변환 후보다.",
        fill=GREEN,
    )
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("※ 확정 동일성과 가중치 기반 파생 후보를 구분해 표시함")
    run.italic = True
    run.font.color.rgb = RGBColor(127, 127, 127)
    document.add_page_break()


def build_document(output_path: Path, pipeline_path: Path, tree_path: Path) -> None:
    """
    목적: 분석 결론·용어·알고리즘·MotherTree를 포함한 최종 DOCX를 조립한다.
    입력: DOCX 출력 경로와 삽입할 파이프라인·계보 PNG 경로.
    처리: 스타일, 표지, 표, 본문, 이미지, 근거 수준, Q&A를 순서대로 구성한다.
    반환/부작용: 반환값 없이 output_path에 DOCX를 저장한다.
    """

    document = Document()
    set_document_font(document)
    set_page_layout(document)
    add_header_footer(document)
    document.core_properties.title = "Qwen3.5 10개 모델 가중치 계보 분석"
    document.core_properties.subject = "Weight distance, kurtosis, CKA, delta SVD, MotherTree"
    document.core_properties.author = "LLM Weight Lineage Research"
    document.core_properties.keywords = "Qwen3.5, lineage, MotherTree, weight distance, kurtosis, CKA, SVD"
    add_title_page(document)

    document.add_heading("1. 경영·팀 요약", level=1)
    add_callout(
        document,
        "무엇이 같고 무엇이 다른가",
        "Qwen 0.8B Instruct와 unsloth 0.8B는 완전히 같다. Qwen 2B Instruct, unsloth 2B, "
        "hamishivi 2B도 완전히 같다. CloudGoat와 두 Huihui는 Instruct에 가장 가깝지만, "
        "CloudGoat는 넓게 조정됐고 Huihui는 O/down만 한 방향으로 정밀하게 바뀌었다.",
        fill=GREEN,
    )
    add_table(
        document,
        ["구분", "가장 비슷한 모델", "차이", "판정"],
        [
            ("unsloth 0.8B", "Qwen Instruct 0.8B", "Weight 차이 없음", "확정"),
            ("CloudGoat 0.8B", "Qwen Instruct 0.8B", "Q/K/V/O 일부와 MLP 전반 변경", "강한 후보"),
            ("Huihui 0.8B", "Qwen Instruct 0.8B", "O/down의 사실상 rank-1 변경", "강한 후보"),
            ("unsloth 2B", "Qwen Instruct 2B", "Weight 차이 없음", "확정"),
            ("hamishivi 2B", "Qwen Instruct 2B", "Weight 차이 없음", "확정"),
            ("Huihui 2B", "Qwen Instruct 2B", "O/down의 사실상 rank-1 변경", "강한 후보"),
        ],
        widths_cm=[3.1, 4.2, 6.7, 2.2],
    )
    add_bullets(
        document,
        [
            "공식 Base와 Instruct는 core 전체에 걸쳐 차이가 나는 별도 anchor다.",
            "0.8B와 2B는 shape가 다르므로 서로 직접 원소 거리를 비교하지 않았다.",
            "파생 후보 방향은 weight 증거에 가장 잘 맞는 관계이며 제작 시간 순서의 확정은 아니다.",
        ],
    )

    document.add_heading("2. 왜 10개가 7개가 됐는가", level=1)
    document.add_paragraph(
        "10개는 다운로드한 모델 저장소·파일 수이고, 7개는 서로 다른 weight 내용의 수다. "
        "세 모델을 분석에서 삭제한 것이 아니라, SHA-256이 같은 파일을 하나의 동일-weight 그룹으로 묶어 "
        "거리 0인 계산을 반복하지 않았다. 계보도와 최종 판정표에는 10개 모두 남아 있다."
    )
    add_table(
        document,
        ["Family", "다운로드한 배포본", "SHA 동일 그룹", "고유 weight"],
        [
            ("0.8B", "Base, Instruct, unsloth, CloudGoat, Huihui", "Instruct = unsloth", "5 - 1 = 4"),
            ("2B", "Base, Instruct, unsloth, hamishivi, Huihui", "Instruct = unsloth = hamishivi", "5 - 2 = 3"),
            ("합계", "배포본 10개", "중복 weight 3개", "4 + 3 = 7"),
        ],
        widths_cm=[2.0, 6.7, 5.0, 2.7],
        font_size=9.0,
    )
    add_callout(
        document,
        "SHA 동일성의 범위",
        "다운로드한 safetensors weight 파일 전체 byte가 같다는 뜻이다. 저장소 설명, tokenizer, "
        "라이선스 표기, 배포 주체까지 동일하다는 의미는 아니다.",
        fill=YELLOW,
    )

    document.add_heading("3. MotherTree: 10개 전체 계보 후보도", level=1)
    document.add_paragraph(
        "굵은 녹색 선은 SHA-256 동일 관계로 확정이다. 주황색 점선은 weight distance와 CKA/SVD가 "
        "지지하는 파생 후보다. 회색 점선은 공식 anchor와 거리로 둔 참고 방향으로, 실제 제작 시점은 미확정이다."
    )
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(tree_path), width=Cm(17.0))
    caption = document.add_paragraph("그림 1. Qwen3.5 10개 모델 MotherTree 후보도")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True
    caption.runs[0].font.size = Pt(9)

    document.add_heading("4. 10개 모델 전체 판정표", level=1)
    add_table(
        document,
        ["#", "모델", "기준", "관계 판정", "핵심 증거", "수준"],
        MODEL_ROWS,
        widths_cm=[0.7, 4.2, 2.5, 3.6, 4.4, 1.5],
        font_size=7.2,
    )
    add_callout(
        document,
        "전체 분류",
        "공식 Base/Instruct 기준점 4개 + 공식 Instruct exact mirror 3개 + Instruct 기반 파생 후보 3개 = 10개",
        fill=LIGHT_BLUE,
    )

    document.add_page_break()
    document.add_heading("5. 분석 알고리즘은 어떻게 작동했는가", level=1)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(pipeline_path), width=Cm(17.0))
    caption = document.add_paragraph("그림 2. Weight 수집부터 MotherTree 후보 생성까지")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True
    caption.runs[0].font.size = Pt(9)
    add_numbered(
        document,
        [
            "10개 safetensors의 tensor 수, parameter 수와 SHA-256을 확인했다.",
            "SHA가 같은 배포본을 한 대표 weight로 묶어 실질 비교 대상을 7개로 줄였다.",
            "0.8B와 2B family를 분리하고 같은 이름·shape·dtype의 tensor만 정렬했다.",
            "전체 모델(all)과 언어 핵심 layer(language_core)를 분리해 불변 영역의 희석을 막았다.",
            "Weight distance와 cosine으로 가까운 후보를, kurtosis로 분포 형태 변화를 조사했다.",
            "CloudGoat의 광범위 변경에는 CKA, Huihui의 국소 변경에는 delta SVD를 적용했다.",
            "Weight distance를 edge 비용, CKA/SVD를 edge 유형으로 사용해 MotherTree 후보를 정리했다.",
        ],
    )
    add_table(
        document,
        ["검증 항목", "결과"],
        [
            ("입력", "10개, 약 29.35GB"),
            ("Tensor 통계", "5,600행, 오류 0"),
            ("고유 모델", "7개"),
            ("Family 내부 pair", "9개"),
            ("Weight-distance tensor pair", "4,824행, shape/dtype skip 0"),
            ("고급 분석", "CKA 96행 + delta SVD 100행"),
            ("재현성", "동일 seed 재실행 결과 일치"),
            ("NAS 회귀 테스트", "13개 통과"),
        ],
        widths_cm=[6.0, 10.2],
    )

    document.add_page_break()
    document.add_heading("6. 전문용어와 측정 원리", level=1)
    document.add_paragraph(
        "거리 하나로 모든 관계를 판단하지 않았다. 각 지표는 서로 다른 질문에 답하므로 독립 증거로 유지했다."
    )
    add_metric_section(
        document,
        "6.1 SHA-256 — 파일 전체의 주민등록번호",
        "파일의 모든 byte를 하나의 긴 식별값으로 만든다. 값이 같으면 이번에 받은 파일 내용 전체가 같다.",
        "암호학적 해시 함수로 입력 전체를 고정 길이 digest로 압축한다. 한 byte만 달라도 일반적으로 digest가 달라진다.",
        "완전 복제본을 pairwise 계산 전에 제거하고 확정 동일 edge를 만든다.",
        "Weight 파일 동일 판정이지 저장소 설명·tokenizer·라이선스 전체의 동일 판정은 아니다.",
    )
    add_metric_section(
        document,
        "6.2 Weight distance — 같은 위치의 다이얼 이동량",
        "두 모델의 같은 기계 설정표에서 같은 위치 숫자가 얼마나 움직였는지 재는 자다.",
        "symmetric L2 = ||A-B||F / sqrt(||A||F² + ||B||F²). 분자로 전체 원소 차이, 분모로 원래 weight 크기를 정규화한다. 0이면 동일하다.",
        "같은 family에서 가장 가까운 모델 후보를 찾고 MotherTree edge 비용으로 사용한다.",
        "거리는 대칭이라 부모→자식 시간 방향을 단독으로 알려주지 않는다. 큰 tensor 영향과 tensor 수 영향을 all/core·median/p95로 나눠 확인해야 한다.",
    )
    add_metric_section(
        document,
        "6.3 Cosine distance — 크기보다 방향이 같은가",
        "두 weight 배열을 긴 화살표로 보고 방향 차이를 잰다.",
        "1 - <A,B> / (||A||F ||B||F)로 계산한다. 패턴 방향이 같으면 숫자 크기가 조금 달라도 작다.",
        "L2와 함께 사용해 단순 scale 변화와 방향 변화를 보조적으로 구분한다.",
        "Cosine이 작아도 실제 원소 이동량인 L2는 클 수 있다.",
    )
    add_metric_section(
        document,
        "6.4 Kurtosis — 숫자 분포 꼬리의 뾰족함",
        "평균이 같은 자갈 자루 두 개에서 유난히 무거운 돌이 얼마나 섞였는지 보는 지표다.",
        "평균에서 멀리 떨어진 값의 4차 모멘트를 분산으로 정규화해 분포의 꼬리와 극단값 민감도를 나타낸다.",
        "Weight가 얼마나 움직였는지가 아니라 분포 모양이 달라졌는지 별도 evidence channel로 확인한다.",
        "일부 tensor만 변하면 전체 median이 0이 될 수 있다. Huihui 탐지에는 p95, 변경 tensor 수, module별 분석이 필요했다.",
    )
    add_metric_section(
        document,
        "6.5 CKA — 내부 관계 구조가 유지됐는가",
        "학생들의 키가 모두 조금 변해도 순서와 학생 간 관계가 유지되면 높은 점수를 주는 방법이다.",
        "Centered Kernel Alignment는 두 matrix가 만드는 행·특징 사이의 중심화된 관계 구조 유사도를 측정한다. 1에 가까울수록 geometry가 유사하다.",
        "광범위하게 바뀐 CloudGoat weight가 원래 Instruct의 구조를 유지하는지 최대 2,048행 결정론적 표본으로 검사했다.",
        "CKA가 높다는 것은 변화량이 작다는 뜻이 아니다. L2가 커도 관계 geometry는 유지될 수 있다.",
    )
    add_metric_section(
        document,
        "6.6 Delta와 SVD — 변화가 몇 개 방향인가",
        "원본에서 후보로 가면서 더해지거나 빠진 변화만 분리하고, 그 변화를 중요한 화살표 순서로 나눈다.",
        "Delta = W후보 - W기준. SVD는 Delta를 singular value와 좌·우 방향으로 분해한다. 첫 singular value 에너지가 대부분이면 rank-1에 가깝다.",
        "Huihui O/down 변화가 복잡한 재학습인지, 반복 가능한 단일 방향 변환인지 판별했다.",
        "최종 delta가 rank-1이라는 사실만으로 LoRA나 특정 도구 사용을 확정할 수 없다.",
    )
    add_metric_section(
        document,
        "6.7 Update alignment — 두 학습 변화 방향이 같은가",
        "Base→Instruct 변화 화살표와 Instruct→후보 변화 화살표가 같은 방향인지 비교한다.",
        "+1이면 같은 방향, -1이면 반대 방향, 0이면 거의 직교한다. 대응하는 delta를 cosine으로 비교했다.",
        "Huihui가 Base→Instruct 학습을 단순 연장하거나 되돌린 것인지 확인했다.",
        "직교는 별도 변화 방향이라는 뜻이며 제작 시점이나 학습 목적의 증명은 아니다.",
    )

    document.add_page_break()
    document.add_heading("7. CloudGoat: 넓게 바뀌었지만 구조는 유지", level=1)
    add_callout(
        document,
        "분류",
        "Instruct-nearest, geometry-preserving substantial core fine-tune candidate",
        fill=YELLOW,
    )
    add_table(
        document,
        ["항목", "결과", "의미"],
        [
            ("최근접 모델", "Qwen 0.8B Instruct", "Base보다 Instruct 계열에 가까움"),
            ("Language-core 거리", "0.042149", "후보 중 변화량이 큼"),
            ("Base→Instruct 대비", "89.12%", "가벼운 몇 군데 수정으로 보기 어려움"),
            ("변경 범위", "Core 192개 중 96개", "Q/K/V/O 일부 + MLP 전반"),
            ("CKA", "모두 0.991 이상", "관계 geometry는 대체로 유지"),
            ("가장 큰 변화", "MLP-up", "CKA 중앙값 0.992794, L2도 가장 큼"),
            ("L2↔CKA distance 상관", "0.92799", "원소 변화와 geometry 변화가 함께 증가"),
            ("Kurtosis↔CKA distance 상관", "-0.00356", "서로 다른 정보를 제공"),
        ],
        widths_cm=[4.0, 4.1, 8.1],
        font_size=8.8,
    )
    document.add_paragraph(
        "해석: CloudGoat는 많은 core 다이얼을 움직였지만 기존 matrix의 관계 구조를 무너뜨리지는 않았다. "
        "따라서 Instruct에 가까운 구조 보존형 광범위 조정 후보로 보는 것이 가장 타당하다. 이름의 JP-Tuned는 "
        "제작자 명칭이며 실제 일본어 성능·학습 데이터는 행동 평가 없이는 확정할 수 없다."
    )

    document.add_heading("8. Huihui: O/down에 찍힌 rank-1 변환", level=1)
    add_callout(
        document,
        "분류",
        "Instruct-nearest, targeted rank-1 projection transformation candidate",
        fill=ORANGE,
    )
    add_table(
        document,
        ["항목", "0.8B", "2B", "의미"],
        [
            ("최근접 모델", "Qwen Instruct 0.8B", "Qwen Instruct 2B", "각 family의 Instruct 계열"),
            ("Language-core 거리", "0.011532", "0.009362", "CloudGoat보다 Instruct에 훨씬 가까움"),
            ("Base→Instruct 대비", "24.38%", "16.55%", "전체 core 변화는 제한적"),
            ("변경 위치", "O/down", "O/down", "Q/K/V와 gate/up은 core에서 동일"),
            ("고급 분석 tensor", "50개", "50개", "합계 100개"),
            ("Top-1 energy 평균", "99.7684%", "99.5593%", "사실상 rank-1"),
            ("Top-1 energy 최솟값", "99.7129%", "99.4660%", "100개 모두 압도적 단일 방향"),
            ("Base→Instruct alignment", "거의 0", "거의 0", "기존 학습과 거의 직교하는 별도 방향"),
        ],
        widths_cm=[3.5, 3.4, 3.4, 6.0],
        font_size=8.4,
    )
    document.add_paragraph(
        "해석: 0.8B와 2B에서 같은 O/down 위치와 같은 rank-1 구조가 반복됐다. 전체를 여러 방향으로 "
        "재학습한 모습보다 특정 projection에 공통된 구조적 변환을 적용한 fingerprint에 가깝다. 그러나 "
        "rank-1만으로 LoRA 사용을 단정할 수 없고, abliterated 기능 효과는 별도의 모델 응답 평가가 필요하다."
    )

    document.add_heading("9. 세 모델 유형을 한눈에 비교", level=1)
    add_table(
        document,
        ["유형", "대표 모델", "변경 범위", "핵심 지표", "해석"],
        [
            ("Exact mirror", "unsloth·hamishivi", "없음", "SHA 동일, 거리 0", "공식 Instruct weight와 완전 동일"),
            ("광범위 조정", "CloudGoat 0.8B", "Q/K/V/O 일부 + MLP 전반", "거리 큼, CKA 높음", "많이 바뀌었지만 geometry 보존"),
            ("표적 변환", "Huihui 0.8B/2B", "O/down 집중", "거리 작음, delta rank-1", "한 방향 projection 변환"),
            ("공식 기준점", "Base·Instruct", "Core 전반 차이", "family 기준 거리", "후보 관계 판단 anchor"),
        ],
        widths_cm=[2.8, 3.6, 4.2, 3.2, 3.3],
        font_size=8.5,
    )

    document.add_heading("10. MotherTree의 원리와 해석 범위", level=1)
    document.add_paragraph(
        "MotherTree는 모델을 node로, 모델 사이 관계를 edge로 표현한다. 이번 단계에서는 SHA 동일 그룹을 먼저 "
        "묶고, 같은 architecture family 안에서 weight distance가 작은 관계를 후보 edge로 선택했다. CKA와 SVD는 "
        "edge가 광범위 구조 보존 조정인지, 국소 저랭크 변환인지 유형을 설명한다."
    )
    add_table(
        document,
        ["도식 요소", "사용 증거", "확실성"],
        [
            ("굵은 녹색 동일선", "SHA-256 전체 파일 동일", "확정"),
            ("주황색 점선", "최근접 weight + 변경 패턴 + CKA/SVD", "강한 후보"),
            ("회색 점선 방향", "공식 Base/Instruct anchor + 거리", "미확정"),
        ],
        widths_cm=[4.5, 8.0, 3.6],
    )
    add_callout(
        document,
        "왜 방향을 확정하지 않는가",
        "Weight distance는 A↔B가 얼마나 가까운지를 말하지만 A가 먼저인지 B가 먼저인지는 말하지 않는다. "
        "공개 시점, base_model 선언, 학습 기록 같은 외부 provenance가 있어야 부모→자식 방향을 강화할 수 있다.",
        fill=YELLOW,
    )

    document.add_heading("11. 핵심 시사점", level=1)
    add_numbered(
        document,
        [
            "저장소 수와 실질 weight 수는 다르다. 10개 중 3개는 공식 Instruct와 완전히 같았다.",
            "전체 평균 하나로 표적 수정을 판단하면 안 된다. Huihui는 전체 kurtosis 중앙값에서 변화가 숨었다.",
            "Q/K만으로는 불충분하다. Huihui 탐지에는 attention O와 MLP down이 필수였다.",
            "거리와 변화 유형은 다른 질문이다. 거리는 가까운 모델을, CKA/SVD는 변화 방식을 설명한다.",
            "확정 동일성, 유사성, 방향성, 기능 효과는 서로 다른 증거 수준으로 관리해야 한다.",
        ],
    )

    document.add_heading("12. 한계와 다음 단계", level=1)
    add_table(
        document,
        ["현재 확인한 것", "아직 확인하지 못한 것", "다음 검증"],
        [
            ("Weight 파일 exact mirror", "저장소 전체·법적 동일성", "라이선스·부가 파일 검토"),
            ("최근접 weight 후보", "실제 제작 시간 순서", "공개 시점·model card·학습 기록 결합"),
            ("CloudGoat 구조 보존 변경", "일본어 성능·학습 데이터", "행동·벤치마크 평가"),
            ("Huihui O/down rank-1 delta", "사용 도구·abliteration 효과", "singular vector·모델 응답 평가"),
            ("0.8B·2B family 내부 관계", "크기 간 직접 원소 계보", "정규화 subspace signature"),
        ],
        widths_cm=[5.1, 5.3, 5.7],
        font_size=8.5,
    )

    document.add_heading("13. 예상 질문과 짧은 답", level=1)
    qa = [
        ("Weight distance가 가장 작으면 부모인가?", "가까운 부모 후보일 뿐이다. 거리에는 시간 방향이 없다."),
        ("SHA가 같으면 완전히 같은 모델인가?", "이번 safetensors weight 파일은 byte 단위로 같다. 저장소 전체는 별도다."),
        ("CKA 0.99면 거의 수정하지 않은 것인가?", "아니다. 변화량이 아니라 관계 geometry가 유지됐다는 뜻이다."),
        ("Rank-1이면 LoRA인가?", "아니다. 최종 변화가 한 방향이라는 사실만 확인했다."),
        ("이 결과로 도용 여부를 증명할 수 있는가?", "아니다. 기술적 동일성·유사성 증거이며 법적 판단은 별도다."),
    ]
    for question, answer in qa:
        paragraph = document.add_paragraph()
        q_run = paragraph.add_run(f"Q. {question}\n")
        q_run.bold = True
        q_run.font.color.rgb = RGBColor.from_string(NAVY)
        paragraph.add_run(f"A. {answer}")

    document.add_heading("14. 팀원이 기억할 네 문장", level=1)
    add_callout(
        document,
        "최종 요약",
        "① 이름이 달라도 SHA가 같으면 weight는 같다.\n"
        "② Weight distance는 변화량, kurtosis는 분포 모양, CKA는 관계 geometry를 잰다.\n"
        "③ SVD는 변화가 한 방향인지 여러 방향인지 보여준다.\n"
        "④ 현재 MotherTree는 강한 가중치 계보 후보도이며 제작 이력을 확정한 법적 족보는 아니다.",
        fill=GREEN,
    )

    document.add_heading("15. 근거 자료", level=1)
    add_bullets(
        document,
        [
            "qwen35_first10_analysis_report.md — 10개 모델 tensor 통계와 SHA 결과",
            "qwen35_weight_distance_report.md — 7개 고유 모델의 실제 weight distance",
            "qwen35_cka_delta_svd_report.md — CloudGoat CKA와 Huihui delta SVD",
            "algorithm_map.md — 프로젝트 계보 분석 알고리즘 원칙",
        ],
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def validate_docx(output_path: Path, expected_images: int = 2) -> dict[str, int]:
    """
    목적: 생성한 DOCX가 OpenXML 구조와 필수 콘텐츠를 갖췄는지 자동 검증한다.
    입력: DOCX 경로와 최소 삽입 이미지 수.
    처리: zip 무결성, media 수, python-docx 재오픈, 표·문단·모델 행을 검사한다.
    반환/부작용: 문단·표·이미지·모델 행 수 dict를 반환하며 누락 시 예외를 발생시킨다.
    """

    if not output_path.is_file() or output_path.stat().st_size < 50_000:
        raise RuntimeError("DOCX가 없거나 예상보다 지나치게 작습니다.")
    with zipfile.ZipFile(output_path) as archive:
        bad_member = archive.testzip()
        if bad_member:
            raise RuntimeError(f"DOCX ZIP 손상: {bad_member}")
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
        if len(media) < expected_images:
            raise RuntimeError(f"삽입 이미지 부족: {len(media)} < {expected_images}")

    document = Document(output_path)
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    required = [
        "왜 10개가 7개가 됐는가",
        "MotherTree",
        "Weight distance",
        "Kurtosis",
        "CKA",
        "Delta와 SVD",
        "CloudGoat",
        "Huihui",
    ]
    missing = [token for token in required if token not in full_text]
    if missing:
        raise RuntimeError(f"필수 문구 누락: {missing}")
    model_row_count = sum(1 for row in MODEL_ROWS if row[1] in full_text or any(row[1] in cell.text for table in document.tables for table_row in table.rows for cell in table_row.cells))
    if model_row_count != 10:
        raise RuntimeError(f"10개 모델 판정표 누락: {model_row_count}/10")
    return {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "images": len(document.inline_shapes),
        "model_rows": model_row_count,
    }


def main() -> None:
    """
    목적: 도식 두 개와 최종 DOCX를 순서대로 생성하고 자동 검증 결과를 출력한다.
    입력: parse_args가 반환한 출력 디렉터리와 글꼴 경로.
    처리: assets 생성, DOCX 조립, OpenXML 검증을 수행한다.
    반환/부작용: 반환값 없이 PNG·DOCX 파일을 만들고 검증 통계를 표준 출력한다.
    """

    args = parse_args()
    output_dir: Path = args.output_dir.resolve()
    font_path: Path = args.font.resolve()
    bold_path: Path = (args.bold_font or args.font).resolve()
    assets_dir = output_dir / "assets"
    pipeline_path = assets_dir / "qwen35_analysis_pipeline.png"
    tree_path = assets_dir / "qwen35_mothertree.png"
    docx_path = output_dir / "Qwen3.5_10-model_lineage_team_report.docx"

    create_pipeline_image(pipeline_path, font_path, bold_path)
    create_mothertree_image(tree_path, font_path, bold_path)
    build_document(docx_path, pipeline_path, tree_path)
    stats = validate_docx(docx_path)
    print(f"DOCX={docx_path}")
    print(f"PIPELINE={pipeline_path}")
    print(f"MOTHERTREE={tree_path}")
    print("VALIDATION=" + ",".join(f"{key}:{value}" for key, value in stats.items()))


if __name__ == "__main__":
    main()
